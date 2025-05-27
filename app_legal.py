import streamlit as st
import datetime
from datetime import datetime
import plotly.graph_objects as go
import re
from docx import Document
from io import BytesIO
from openai import OpenAI
import os

# Función para validar y obtener la API key
def get_openai_api_key():
    api_key = os.getenv('OPENAI_API_KEY')
    if not api_key:
        st.error("⚠️ No se encontró la API key de OpenAI. Por favor, asegúrese de que la variable de entorno OPENAI_API_KEY está configurada correctamente.")
        st.info("Para configurar la API key permanentemente, agregue la siguiente línea a su archivo ~/.zshrc:")
        st.code('export OPENAI_API_KEY="su-api-key-aquí"')
        st.info("Luego, reinicie su terminal o ejecute: source ~/.zshrc")
        return None
    return api_key

# Configurar el cliente de OpenAI con validación
api_key = get_openai_api_key()
if api_key:
    try:
        client = OpenAI(api_key=api_key)
    except Exception as e:
        st.error(f"⚠️ Error al configurar el cliente de OpenAI: {str(e)}")
        client = None
else:
    client = None

# Set page config
st.set_page_config(
    page_title="Asistente Legal Peruano",
    page_icon="⚖️",
    layout="centered"
)

# Custom CSS with elegant earth tones and enhanced styling
st.markdown("""
    <style>
    .main {
        background: linear-gradient(135deg, #f5f1e6 0%, #e6d5c3 100%);
    }
    .stButton>button {
        background: linear-gradient(45deg, #8B4513, #A0522D);
        color: #f5f1e6;
        border-radius: 8px;
        padding: 12px 30px;
        border: none;
        transition: all 0.3s ease;
        font-family: 'Times New Roman', serif;
        font-weight: 500;
    }
    .stButton>button:hover {
        background: linear-gradient(45deg, #A0522D, #8B4513);
        transform: translateY(-2px);
        box-shadow: 0 4px 8px rgba(139, 69, 19, 0.2);
    }
    h1, h2, h3 {
        color: #5C4033;
        font-family: 'Times New Roman', serif;
        text-align: center;
        text-shadow: 1px 1px 2px rgba(0,0,0,0.1);
        margin: 0;
        padding: 10px 0;
    }
    .stSelectbox, .stTextInput, .stTextArea {
        background: rgba(255, 255, 255, 0.95);
        border-radius: 8px;
        padding: 12px;
        border: 1px solid #D2B48C;
        font-family: 'Times New Roman', serif;
    }
    .stMarkdown {
        background: rgba(255, 255, 255, 0.85);
        border-radius: 8px;
        padding: 20px;
        margin: 15px 0;
        border: 1px solid #D2B48C;
    }
    .success-box {
        background: linear-gradient(45deg, #8B7355, #A67B5B);
        color: #f5f1e6;
        padding: 20px;
        border-radius: 8px;
        margin: 15px 0;
        border: 1px solid #D2B48C;
    }
    .warning-box {
        background: linear-gradient(45deg, #A67B5B, #8B7355);
        color: #f5f1e6;
        padding: 20px;
        border-radius: 8px;
        margin: 15px 0;
        border: 1px solid #D2B48C;
    }
    .info-box {
        background: linear-gradient(45deg, #8B7355, #A67B5B);
        color: #f5f1e6;
        padding: 20px;
        border-radius: 8px;
        margin: 15px 0;
        border: 1px solid #D2B48C;
    }
    .sidebar .sidebar-content {
        background: linear-gradient(135deg, #f5f1e6 0%, #e6d5c3 100%);
    }
    .stForm {
        background: rgba(255, 255, 255, 0.95);
        border-radius: 8px;
        padding: 20px;
        margin: 15px 0;
        border: 1px solid #D2B48C;
    }
    .stRadio > div {
        background: rgba(255, 255, 255, 0.95);
        border-radius: 8px;
        padding: 10px;
        border: 1px solid #D2B48C;
    }
    .stCheckbox > div {
        background: rgba(255, 255, 255, 0.95);
        border-radius: 8px;
        padding: 10px;
        border: 1px solid #D2B48C;
    }
    .stProgress > div > div > div {
        background-color: #8B4513;
    }
    .stProgress > div > div > div > div {
        background-color: #A0522D;
    }
    .action-card {
        background: linear-gradient(135deg, #f5f1e6 0%, #e6d5c3 100%);
        padding: 20px;
        border-radius: 10px;
        text-align: center;
        width: 30%;
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
        margin: 10px;
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        min-height: 200px;
    }
    .action-card img {
        width: 50px;
        height: 50px;
        margin-bottom: 15px;
    }
    .action-card h3 {
        color: #5C4033;
        font-family: 'Times New Roman', serif;
        margin: 10px 0;
        font-size: 1.2em;
    }
    .action-card p {
        font-size: 0.9em;
        margin: 10px 0;
        color: #5C4033;
    }
    .section-title {
        background: linear-gradient(45deg, #8B7355, #A67B5B);
        color: #f5f1e6;
        padding: 15px;
        border-radius: 8px;
        margin: 20px 0;
        text-align: center;
        font-family: 'Times New Roman', serif;
        font-size: 1.5em;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    </style>
    """, unsafe_allow_html=True)

# Definir las etapas del proceso judicial peruano
etapas_proceso = {
    "Civil": [
        "Presentación de la Demanda",
        "Admisión de la Demanda",
        "Notificación al Demandado",
        "Contestación de la Demanda",
        "Audiencia Única",
        "Sentencia",
        "Ejecución de Sentencia"
    ],
    "Laboral": [
        "Presentación de la Demanda",
        "Admisión de la Demanda",
        "Notificación al Demandado",
        "Contestación de la Demanda",
        "Audiencia Única",
        "Sentencia",
        "Ejecución de Sentencia"
    ],
    "Penal": [
        "Denuncia Penal",
        "Investigación Preliminar",
        "Formalización de la Investigación",
        "Etapa Intermedia",
        "Juicio Oral",
        "Sentencia",
        "Ejecución de la Pena"
    ]
}

# Definir las acciones recomendadas según el tipo de documento
acciones_recomendadas = {
    "Demanda Civil": {
        "Presentación": [
            "Preparar copias del documento para notificación",
            "Verificar el pago de tasas judiciales",
            "Preparar los anexos correspondientes"
        ],
        "Admisión": [
            "Revisar la resolución de admisión",
            "Verificar los plazos para notificación",
            "Preparar el oficio de notificación"
        ],
        "Notificación": [
            "Verificar la constancia de notificación",
            "Preparar la contestación de la demanda",
            "Reunir pruebas para la contestación"
        ],
        "Contestación": [
            "Preparar la audiencia única",
            "Organizar las pruebas a presentar",
            "Preparar los alegatos"
        ]
    },
    "Demanda Laboral": {
        "Presentación": [
            "Verificar la competencia del juzgado laboral",
            "Preparar los documentos laborales",
            "Calcular los montos reclamados"
        ],
        "Admisión": [
            "Revisar la resolución de admisión",
            "Verificar los plazos para notificación",
            "Preparar el oficio de notificación"
        ],
        "Notificación": [
            "Verificar la constancia de notificación",
            "Preparar la contestación de la demanda",
            "Reunir pruebas para la contestación"
        ],
        "Contestación": [
            "Preparar la audiencia única",
            "Organizar las pruebas a presentar",
            "Preparar los alegatos"
        ]
    },
    "Demanda Penal": {
        "Denuncia": [
            "Verificar la competencia del juzgado penal",
            "Preparar las pruebas iniciales",
            "Verificar los plazos de prescripción"
        ],
        "Investigación": [
            "Preparar la declaración del denunciante",
            "Reunir pruebas adicionales",
            "Verificar los plazos de investigación"
        ],
        "Formalización": [
            "Preparar la acusación",
            "Organizar las pruebas",
            "Verificar los plazos procesales"
        ]
    }
}

# Definir tipos de procesos civiles especializados
tipos_procesos_civiles = {
    "Proceso de Familia": {
        "artículos": ["546", "547", "548", "549", "550"],
        "características": ["alimentos", "tenencia", "régimen de visitas", "divorcio", "violencia familiar"],
        "etapas": ["Demanda", "Audiencia Única", "Sentencia"],
        "jurisprudencia": {
            "Pleno Casatorio": [
                "CASACIÓN N° 1-2019/CIJ-116: Sobre alimentos",
                "CASACIÓN N° 2-2020/CIJ-117: Sobre tenencia"
            ],
            "doctrina": [
                "Doctrina sobre alimentos según el Código Civil",
                "Interpretación de la tenencia compartida"
            ]
        }
    },
    "Proceso de Derechos Reales": {
        "artículos": ["546", "547", "548"],
        "características": ["propiedad", "posesión", "usucapión", "servidumbre"],
        "etapas": ["Demanda", "Contestación", "Audiencia Única", "Sentencia"],
        "jurisprudencia": {
            "Pleno Casatorio": [
                "CASACIÓN N° 3-2018/CIJ-118: Sobre usucapión",
                "CASACIÓN N° 4-2021/CIJ-119: Sobre servidumbre"
            ],
            "doctrina": [
                "Doctrina sobre propiedad según el Código Civil",
                "Interpretación de la posesión"
            ]
        }
    },
    "Proceso de Garantías": {
        "artículos": ["546", "547", "548"],
        "características": ["medida cautelar", "garantía", "prenda", "hipoteca"],
        "etapas": ["Solicitud", "Resolución", "Ejecución"],
        "jurisprudencia": {
            "Pleno Casatorio": [
                "CASACIÓN N° 5-2017/CIJ-120: Sobre medidas cautelares",
                "CASACIÓN N° 6-2022/CIJ-121: Sobre garantías"
            ],
            "doctrina": [
                "Doctrina sobre medidas cautelares",
                "Interpretación de garantías"
            ]
        }
    },
    "Proceso de Conocimiento": {
        "artículos": ["506", "507", "508", "509", "510"],
        "características": ["pretensión de condena", "ejecución de obligación", "cumplimiento de contrato"],
        "etapas": ["Demanda", "Contestación", "Audiencia Única", "Sentencia"],
        "jurisprudencia": {
            "Pleno Casatorio": [
                "CASACIÓN N° 7-2016/CIJ-122: Sobre obligaciones",
                "CASACIÓN N° 8-2023/CIJ-123: Sobre contratos"
            ],
            "doctrina": [
                "Doctrina sobre obligaciones",
                "Interpretación de contratos"
            ]
        }
    },
    "Proceso Abreviado": {
        "artículos": ["546", "547", "548"],
        "características": ["cuantía determinada", "prueba documental", "plazo reducido"],
        "etapas": ["Demanda", "Contestación", "Audiencia Única", "Sentencia"],
        "jurisprudencia": {
            "Pleno Casatorio": [
                "CASACIÓN N° 9-2015/CIJ-124: Sobre proceso abreviado",
                "CASACIÓN N° 10-2024/CIJ-125: Sobre plazos procesales"
            ],
            "doctrina": [
                "Doctrina sobre proceso abreviado",
                "Interpretación de plazos"
            ]
        }
    }
}

# Definir acciones finales según el tipo de proceso
acciones_finales = {
    "Proceso de Conocimiento": [
        "Presentar apelación dentro de los 5 días hábiles",
        "Solicitar ejecución de sentencia si es condenatoria",
        "Preparar recursos de casación si corresponde"
    ],
    "Proceso Abreviado": [
        "Presentar apelación dentro de los 3 días hábiles",
        "Solicitar ejecución de sentencia si es condenatoria",
        "Verificar plazos para recursos extraordinarios"
    ],
    "Proceso Sumarísimo": [
        "Presentar apelación dentro de los 2 días hábiles",
        "Solicitar ejecución inmediata si es condenatoria",
        "Verificar plazos para recursos"
    ],
    "Proceso de Ejecución": [
        "Presentar oposición a la ejecución si corresponde",
        "Solicitar levantamiento de medidas cautelares",
        "Preparar recursos contra la sentencia"
    ],
    "Proceso Cautelar": [
        "Solicitar levantamiento de medida cautelar",
        "Presentar garantía si corresponde",
        "Preparar demanda principal"
    ]
}

# Definir acciones específicas según el tipo de proceso
acciones_especificas = {
    "Proceso de Familia": {
        "Demanda": [
            "Verificar competencia del juzgado de familia",
            "Anexar partidas de nacimiento",
            "Incluir prueba de ingresos para alimentos"
        ],
        "Audiencia": [
            "Preparar conciliación",
            "Organizar pruebas testimoniales",
            "Preparar informe psicológico si corresponde"
        ],
        "Sentencia": [
            "Verificar plazos de apelación (5 días)",
            "Preparar recursos si corresponde",
            "Solicitar ejecución de sentencia"
        ]
    },
    "Proceso de Derechos Reales": {
        "Demanda": [
            "Anexar título de propiedad",
            "Incluir plano de ubicación",
            "Presentar prueba de posesión"
        ],
        "Contestación": [
            "Preparar oposición fundada",
            "Anexar documentos de propiedad",
            "Solicitar peritaje si corresponde"
        ],
        "Sentencia": [
            "Verificar plazos de apelación",
            "Preparar recursos",
            "Solicitar ejecución de sentencia"
        ]
    },
    "Proceso de Garantías": {
        "Solicitud": [
            "Fundamentar urgencia",
            "Anexar documentos de garantía",
            "Ofrecer contragarantía"
        ],
        "Resolución": [
            "Verificar cumplimiento de requisitos",
            "Preparar oposición si corresponde",
            "Solicitar modificación si es necesario"
        ],
        "Ejecución": [
            "Verificar plazos de ejecución",
            "Preparar recursos",
            "Solicitar levantamiento si corresponde"
        ]
    }
}

# Definir diccionario jurídico por tipo de proceso
diccionario_juridico = {
    "Civil": {
        "Presentación": {
            "palabras_clave": [
                "presento", "interpongo", "demanda", "solicito", "vengo a interponer",
                "a ud. respetuosamente digo", "fundamentos de hecho", "fundamentos de derecho",
                "petitorio", "pretensión", "derecho subjetivo", "interés legítimo",
                "competencia", "cuantía", "valor de la pretensión"
            ],
            "documentos": [
                "demanda", "escrito de demanda", "solicitud inicial", "petitorio",
                "poder", "documento de identidad", "documentos de respaldo",
                "recibo de pago de tasas judiciales", "declaración jurada"
            ],
            "fechas_clave": [
                "fecha de presentación", "fecha de interposición", "fecha de ingreso",
                "fecha de vencimiento", "plazo de caducidad"
            ]
        },
        "Admisión": {
            "palabras_clave": [
                "admito", "admisión", "expediente", "número de expediente",
                "resolución de admisión", "auto admisorio", "competencia",
                "cuantía", "plazo", "requisitos de admisibilidad"
            ],
            "documentos": [
                "resolución de admisión", "auto admisorio",
                "resolución que admite la demanda", "oficio de devolución",
                "resolución que declara inadmisible"
            ],
            "fechas_clave": [
                "fecha de admisión", "fecha de expediente", "plazo para subsanar",
                "plazo para notificar"
            ]
        },
        "Notificación": {
            "palabras_clave": [
                "notificar", "notificación", "citación", "emplazamiento",
                "constancia de notificación", "oficio de notificación",
                "domicilio procesal", "domicilio real", "edictos",
                "publicación", "plazo para contestar"
            ],
            "documentos": [
                "oficio de notificación", "constancia de notificación",
                "certificado de notificación", "edicto", "publicación",
                "constancia de emplazamiento"
            ],
            "fechas_clave": [
                "fecha de notificación", "fecha de citación", "plazo para contestar",
                "fecha de emplazamiento"
            ]
        },
        "Contestación": {
            "palabras_clave": [
                "contesto", "contestación", "oposición", "excepciones",
                "defensa", "alegatos de contestación", "excepción de incompetencia",
                "excepción de litispendencia", "excepción de cosa juzgada",
                "excepción de prescripción", "reconvención"
            ],
            "documentos": [
                "escrito de contestación", "contestación de demanda",
                "excepciones y defensa previa", "reconvención",
                "pruebas de contestación"
            ],
            "fechas_clave": [
                "fecha de contestación", "plazo para contestar",
                "plazo para reconvenir"
            ]
        },
        "Audiencia": {
            "palabras_clave": [
                "audiencia", "vista", "comparecencia", "conciliación",
                "alegatos", "pruebas", "testigos", "peritos",
                "inspección judicial", "reconocimiento de documentos",
                "juramento", "declaración de parte"
            ],
            "documentos": [
                "acta de audiencia", "constancia de audiencia",
                "alegatos orales", "acta de conciliación",
                "constancia de pruebas"
            ],
            "fechas_clave": [
                "fecha de audiencia", "día de la audiencia",
                "plazo para presentar pruebas"
            ]
        },
        "Sentencia": {
            "palabras_clave": [
                "sentencia", "fallo", "resolución", "decisión",
                "condena", "absolución", "fundamentos de hecho",
                "fundamentos de derecho", "dispositivo", "costas y costos",
                "apelación", "recurso de apelación"
            ],
            "documentos": [
                "sentencia", "resolución que pone fin al proceso",
                "fallo", "resolución de apelación"
            ],
            "fechas_clave": [
                "fecha de sentencia", "fecha del fallo",
                "plazo para apelar"
            ]
        }
    },
    "Laboral": {
        "Presentación": {
            "palabras_clave": [
                "presento", "interpongo", "demanda laboral", "solicito",
                "relación laboral", "contrato de trabajo", "despido",
                "cese", "beneficios sociales", "compensación por tiempo de servicios",
                "gratificaciones", "vacaciones"
            ],
            "documentos": [
                "demanda laboral", "boletas de pago", "contrato de trabajo",
                "carta de despido", "constancia de trabajo"
            ],
            "fechas_clave": [
                "fecha de presentación", "fecha de cese",
                "fecha de despido"
            ]
        }
    },
    "Penal": {
        "Denuncia": {
            "palabras_clave": [
                "denuncio", "denuncia penal", "delito", "infracción penal",
                "hecho punible", "tipicidad", "antijuridicidad",
                "culpabilidad", "penalidad", "agravantes", "atenuantes"
            ],
            "documentos": [
                "denuncia penal", "parte policial", "atestado",
                "informe pericial", "declaración de testigos"
            ],
            "fechas_clave": [
                "fecha de denuncia", "fecha del delito",
                "plazo de prescripción"
            ]
        }
    }
}

# Actualizar las plantillas de documentos con modelos peruanos
plantillas_documentos = {
    "Presentación": {
        "Civil": {
            "nombre": "Demanda Civil",
            "plantilla": """
            EXPEDIENTE N° {numero_expediente}
            JUZGADO ESPECIALIZADO CIVIL DE LIMA
            {juzgado}
            
            SEÑOR JUEZ:
            
            {demandante}, identificado con DNI/RUC N° {dni_demandante}, con domicilio procesal en {domicilio_demandante}, a Ud. respetuosamente digo:
            
            Que, vengo a interponer DEMANDA DE CONOCIMIENTO contra {demandado}, identificado con DNI/RUC N° {dni_demandado}, con domicilio en {domicilio_demandado}, por los siguientes fundamentos:
            
            I. HECHOS
            {hechos}
            
            II. FUNDAMENTOS DE DERECHO
            {fundamentos}
            
            III. PETITORIO
            Por lo expuesto, solicito a Su Señoría:
            1. Tener por interpuesta la presente demanda de conocimiento.
            2. Ordenar la notificación de la demanda al demandado.
            3. {petitorio}
            
            IV. ANEXOS
            1. Copia simple del DNI del demandante.
            2. Poder notarial que acredita la representación.
            3. Recibo de pago de tasas judiciales.
            4. {anexos}
            
            Es justicia que solicito.
            
            Lima, {fecha}
            
            _________________________
            {demandante}
            """
        },
        "Laboral": {
            "nombre": "Demanda Laboral",
            "plantilla": """
            EXPEDIENTE N° {numero_expediente}
            JUZGADO ESPECIALIZADO LABORAL DE LIMA
            {juzgado}
            
            SEÑOR JUEZ:
            
            {demandante}, identificado con DNI N° {dni_demandante}, con domicilio procesal en {domicilio_demandante}, a Ud. respetuosamente digo:
            
            Que, vengo a interponer DEMANDA LABORAL contra {demandado}, identificado con RUC N° {dni_demandado}, con domicilio en {domicilio_demandado}, por los siguientes fundamentos:
            
            I. HECHOS
            {hechos}
            
            II. RELACIÓN LABORAL
            - Fecha de inicio: {fecha_inicio}
            - Fecha de cese: {fecha_cese}
            - Último salario: S/. {ultimo_salario}
            - Cargo desempeñado: {cargo}
            
            III. PETITORIO
            Por lo expuesto, solicito a Su Señoría:
            1. Tener por interpuesta la presente demanda laboral.
            2. Ordenar la notificación de la demanda al demandado.
            3. {petitorio}
            
            IV. ANEXOS
            1. Copia simple del DNI del demandante.
            2. Poder notarial que acredita la representación.
            3. Recibo de pago de tasas judiciales.
            4. Boletas de pago.
            5. Constancia de trabajo.
            6. {anexos}
            
            Es justicia que solicito.
            
            Lima, {fecha}
            
            _________________________
            {demandante}
            """
        }
    },
    "Contestación": {
        "Civil": {
            "nombre": "Contestación de Demanda",
            "plantilla": """
            EXPEDIENTE N° {numero_expediente}
            JUZGADO ESPECIALIZADO CIVIL DE LIMA
            {juzgado}
            
            SEÑOR JUEZ:
            
            {demandado}, identificado con DNI/RUC N° {dni_demandado}, con domicilio procesal en {domicilio_demandado}, a Ud. respetuosamente digo:
            
            Que, vengo a contestar la demanda interpuesta por {demandante}, por los siguientes fundamentos:
            
            I. EXCEPCIONES PRELIMINARES
            {excepciones}
            
            II. DEFENSA DE FONDO
            {defensa}
            
            III. RECONVENCIÓN (si corresponde)
            {reconvencion}
            
            IV. ANEXOS
            1. Copia simple del DNI del demandado.
            2. Poder notarial que acredita la representación.
            3. {anexos}
            
            Por lo expuesto, solicito a Su Señoría:
            1. Tener por contestada la presente demanda.
            2. {petitorio}
            
            Es justicia que solicito.
            
            Lima, {fecha}
            
            _________________________
            {demandado}
            """
        }
    },
    "Audiencia": {
        "Civil": {
            "nombre": "Escrito de Alegatos",
            "plantilla": """
            EXPEDIENTE N° {numero_expediente}
            JUZGADO ESPECIALIZADO CIVIL DE LIMA
            {juzgado}
            
            SEÑOR JUEZ:
            
            {parte}, identificado con DNI/RUC N° {dni_parte}, con domicilio procesal en {domicilio_parte}, a Ud. respetuosamente digo:
            
            Que, vengo a presentar mis ALEGATOS ORALES para la audiencia programada para el día {fecha_audiencia}, por los siguientes fundamentos:
            
            I. ALEGATOS
            {alegatos}
            
            II. PRUEBAS
            {pruebas}
            
            III. ANEXOS
            1. Copia simple del DNI.
            2. Poder notarial que acredita la representación.
            3. {anexos}
            
            Por lo expuesto, solicito a Su Señoría:
            1. Tener por presentados los presentes alegatos.
            2. {petitorio}
            
            Es justicia que solicito.
            
            Lima, {fecha}
            
            _________________________
            {parte}
            """
        }
    }
}

def create_process_thermometer(current_stage, total_stages):
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=(current_stage / total_stages) * 100,
        domain={'x': [0, 1], 'y': [0, 1]},
        gauge={
            'axis': {'range': [0, 100]},
            'bar': {'color': "#3498db"},
            'steps': [
                {'range': [0, 100], 'color': "#e0e0e0"}
            ],
            'threshold': {
                'line': {'color': "red", 'width': 4},
                'thickness': 0.75,
                'value': (current_stage / total_stages) * 100
            }
        },
        title={'text': "Progreso del Proceso"}
    ))
    
    fig.update_layout(
        height=300,
        margin=dict(l=20, r=20, t=50, b=20),
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)"
    )
    
    return fig

def detectar_articulos(texto):
    # Patrón para encontrar referencias a artículos
    patron_articulos = r'(?:artículo|art\.)\s*(\d+)(?:\s*(?:y|,)\s*(\d+))?'
    articulos = re.findall(patron_articulos, texto.lower())
    return [art for sublist in articulos for art in sublist if art]

def detectar_tipo_proceso(texto, articulos):
    tipo_proceso = None
    max_coincidencias = 0
    texto = texto.lower()
    
    # Palabras clave específicas para cada tipo de proceso
    palabras_clave = {
        "Proceso de Familia": {
            "terminos": ["hijos", "alimentos", "tenencia", "régimen de visitas", "divorcio", "violencia familiar", 
                        "pensión alimenticia", "patria potestad", "guarda", "custodia", "visitas"],
            "articulos": ["546", "547", "548", "549", "550", "551", "552", "553", "554", "555"]
        },
        "Proceso de Derechos Reales": {
            "terminos": ["propiedad", "posesión", "usucapión", "servidumbre", "hipoteca", "prenda", 
                        "desalojo", "desocupación", "título de propiedad", "plano de ubicación"],
            "articulos": ["546", "547", "548", "549", "550", "551", "552", "553", "554", "555"]
        },
        "Proceso de Garantías": {
            "terminos": ["medida cautelar", "garantía", "prenda", "hipoteca", "embargo", "secuestro", 
                        "inhibición", "anotación", "inscripción", "registro"],
            "articulos": ["546", "547", "548", "549", "550", "551", "552", "553", "554", "555"]
        },
        "Proceso de Conocimiento": {
            "terminos": ["pretensión de condena", "ejecución de obligación", "cumplimiento de contrato", 
                        "daños y perjuicios", "indemnización", "resolución de contrato"],
            "articulos": ["506", "507", "508", "509", "510", "511", "512", "513", "514", "515"]
        }
    }
    
    for tipo, info in palabras_clave.items():
        coincidencias = 0
        # Verificar artículos
        for art in articulos:
            if art in info["articulos"]:
                coincidencias += 3  # Mayor peso a coincidencias de artículos
        
        # Verificar términos específicos
        for termino in info["terminos"]:
            if termino in texto:
                coincidencias += 2  # Peso medio a coincidencias de términos
        
        if coincidencias > max_coincidencias:
            max_coincidencias = coincidencias
            tipo_proceso = tipo
    
    return tipo_proceso, max_coincidencias

def detectar_etapa_proceso(texto, tipo_proceso="Civil"):
    # Obtener el diccionario específico para el tipo de proceso
    etapas_criterios = diccionario_juridico.get(tipo_proceso, diccionario_juridico["Civil"])
    
    # Puntuación para cada etapa
    puntuaciones = {etapa: 0 for etapa in etapas_criterios.keys()}
    
    # Analizar el texto
    texto = texto.lower()
    
    # Palabras clave específicas para cada etapa
    palabras_clave_etapas = {
        "Presentación": {
            "inicio": ["presento", "interpongo", "solicito", "vengo a", "a ud. respetuosamente digo"],
            "contenido": ["fundamentos de hecho", "fundamentos de derecho", "petitorio", "pretensión"],
            "documentos": ["demanda", "escrito de demanda", "solicitud inicial"]
        },
        "Admisión": {
            "inicio": ["admito", "admisión", "expediente", "número de expediente"],
            "contenido": ["resolución de admisión", "auto admisorio", "competencia"],
            "documentos": ["resolución de admisión", "auto admisorio"]
        },
        "Notificación": {
            "inicio": ["notificar", "notificación", "citación", "emplazamiento"],
            "contenido": ["constancia de notificación", "oficio de notificación", "domicilio procesal"],
            "documentos": ["oficio de notificación", "constancia de notificación"]
        },
        "Contestación": {
            "inicio": ["contesto", "contestación", "oposición", "excepciones"],
            "contenido": ["defensa", "alegatos de contestación", "excepción de incompetencia"],
            "documentos": ["escrito de contestación", "contestación de demanda"]
        },
        "Audiencia": {
            "inicio": ["audiencia", "vista", "comparecencia", "conciliación"],
            "contenido": ["alegatos", "pruebas", "testigos", "peritos"],
            "documentos": ["acta de audiencia", "constancia de audiencia"]
        },
        "Sentencia": {
            "inicio": ["sentencia", "fallo", "resolución", "decisión"],
            "contenido": ["condena", "absolución", "fundamentos de hecho", "fundamentos de derecho"],
            "documentos": ["sentencia", "resolución que pone fin al proceso"]
        }
    }
    
    for etapa, criterios in palabras_clave_etapas.items():
        # Verificar palabras de inicio
        for palabra in criterios["inicio"]:
            if palabra in texto:
                puntuaciones[etapa] += 3  # Mayor peso a palabras de inicio
        
        # Verificar contenido específico
        for palabra in criterios["contenido"]:
            if palabra in texto:
                puntuaciones[etapa] += 2  # Peso medio a contenido específico
        
        # Verificar documentos
        for documento in criterios["documentos"]:
            if documento in texto:
                puntuaciones[etapa] += 2  # Peso medio a documentos
    
    # Determinar la etapa con mayor puntuación
    etapa_detectada = max(puntuaciones.items(), key=lambda x: x[1])
    
    # Calcular nivel de confianza basado en la puntuación
    puntuacion_maxima = max(puntuaciones.values())
    total_posible = len(palabras_clave_etapas[etapa_detectada[0]]["inicio"]) * 3 + \
                   len(palabras_clave_etapas[etapa_detectada[0]]["contenido"]) * 2 + \
                   len(palabras_clave_etapas[etapa_detectada[0]]["documentos"]) * 2
    
    porcentaje_confianza = (puntuacion_maxima / total_posible) * 100
    
    if porcentaje_confianza >= 70:
        nivel_confianza = "alto"
    elif porcentaje_confianza >= 40:
        nivel_confianza = "medio"
    else:
        nivel_confianza = "bajo"
    
    return etapa_detectada[0], nivel_confianza, puntuaciones

def obtener_fundamentos_jurisprudencia(casacion):
    # Simulación de fundamentos basados en la jurisprudencia
    fundamentos = {
        "CASACIÓN N° 1-2019/CIJ-116": "La obligación alimentaria debe ser proporcional a las necesidades del alimentista y las posibilidades del alimentante.",
        "CASACIÓN N° 2-2020/CIJ-117": "La tenencia compartida debe primar cuando sea beneficiosa para el menor.",
        "CASACIÓN N° 3-2018/CIJ-118": "La usucapión requiere posesión continua, pacífica y pública.",
        "CASACIÓN N° 4-2021/CIJ-119": "La servidumbre debe ser necesaria y proporcionada."
    }
    return fundamentos.get(casacion, "Fundamentos no disponibles.")

def analizar_documento_con_openai(texto, tipo_proceso):
    # Verificar si el cliente de OpenAI está disponible
    if not client:
        st.error("⚠️ No se puede realizar el análisis porque el cliente de OpenAI no está configurado correctamente.")
        st.info("Por favor, configure la API key de OpenAI siguiendo las instrucciones anteriores.")
        return None

    # Preparar el prompt para OpenAI con la estructura JSON escapada correctamente
    prompt = f"""Por favor, analiza el siguiente documento legal de tipo {tipo_proceso}. 
    Necesito que me proporciones un análisis estructurado en formato JSON con la siguiente información exacta:
    
    1. La etapa del proceso (debe ser una de estas: "Presentación", "Admisión", "Notificación", "Contestación", "Audiencia", "Sentencia")
    2. El nivel de confianza de la detección (debe ser: "alto", "medio" o "bajo")
    3. Los artículos mencionados en el texto
    4. La jurisprudencia relevante mencionada
    5. Los campos faltantes importantes
    6. Recomendaciones para el siguiente paso
    
    IMPORTANTE: Tu respuesta debe ser ÚNICAMENTE un objeto JSON válido con esta estructura exacta:
    {{
        "etapa": "string",
        "nivel_confianza": "string",
        "articulos": ["string"],
        "jurisprudencia": ["string"],
        "campos_faltantes": ["string"],
        "recomendaciones": ["string"]
    }}
    
    No incluyas ningún otro texto o explicación, solo el JSON.
    
    Documento a analizar:
    {texto}
    """

    try:
        # Llamar a la API de OpenAI
        completion = client.chat.completions.create(
            model="gpt-4o-mini",
            store=True,
            messages=[
                {"role": "system", "content": "Eres un asistente legal especializado en análisis de documentos jurídicos. Tus respuestas deben ser siempre en formato JSON válido, sin texto adicional."},
                {"role": "user", "content": prompt}
            ]
        )

        # Obtener la respuesta
        respuesta = completion.choices[0].message.content.strip()
        
        # Asegurarse de que la respuesta comience y termine con llaves
        if not respuesta.startswith('{'): respuesta = '{' + respuesta
        if not respuesta.endswith('}'): respuesta = respuesta + '}'
        
        try:
            # Intentar parsear el JSON
            import json
            analisis = json.loads(respuesta)
            
            # Validar que el JSON tiene la estructura esperada
            campos_requeridos = ["etapa", "nivel_confianza", "articulos", "jurisprudencia", "campos_faltantes", "recomendaciones"]
            for campo in campos_requeridos:
                if campo not in analisis:
                    analisis[campo] = [] if campo in ["articulos", "jurisprudencia", "campos_faltantes", "recomendaciones"] else ""
            
            # Asegurar que los campos de lista sean listas
            for campo in ["articulos", "jurisprudencia", "campos_faltantes", "recomendaciones"]:
                if not isinstance(analisis[campo], list):
                    analisis[campo] = [analisis[campo]] if analisis[campo] else []
            
            # Asegurar que el nivel de confianza sea válido
            if analisis["nivel_confianza"] not in ["alto", "medio", "bajo"]:
                analisis["nivel_confianza"] = "medio"
            
            return analisis
            
        except json.JSONDecodeError as e:
            st.error(f"Error al procesar la respuesta de OpenAI: {str(e)}")
            st.error("Respuesta recibida: " + respuesta)
            
            # Crear un análisis por defecto
            return {
                "etapa": "Presentación",
                "nivel_confianza": "bajo",
                "articulos": [],
                "jurisprudencia": [],
                "campos_faltantes": ["No se pudo determinar"],
                "recomendaciones": ["Se recomienda revisar el documento manualmente"]
            }
            
    except Exception as e:
        st.error(f"Error al comunicarse con OpenAI: {str(e)}")
        return None

def obtener_siguiente_documento(etapa_actual, tipo_proceso):
    # Mapeo de etapas a documentos siguientes
    siguiente_documento = {
        "Presentación": "Contestación",
        "Admisión": "Notificación",
        "Notificación": "Contestación",
        "Contestación": "Audiencia",
        "Audiencia": "Sentencia",
        "Sentencia": "Apelación"
    }
    
    siguiente_etapa = siguiente_documento.get(etapa_actual)
    if siguiente_etapa and siguiente_etapa in plantillas_documentos:
        return plantillas_documentos[siguiente_etapa].get(tipo_proceso)
    return None

def generar_siguiente_documento(etapa_actual, tipo_proceso, datos, documento_original=None, analisis=None):
    # Crear un nuevo documento
    doc = Document()
    
    # Mapeo de documentos siguientes según la etapa actual
    siguiente_documento = {
        "Presentación": {
            "nombre": "Contestación de Demanda",
            "instrucciones": "Este documento debe ser presentado en respuesta a la demanda recibida, dentro del plazo de 15 días hábiles.",
            "plantilla": """
            EXPEDIENTE N° {numero_expediente}
            JUZGADO ESPECIALIZADO CIVIL DE LIMA
            {juzgado}
            
            SEÑOR JUEZ:
            
            {demandado}, identificado con DNI N° {dni_demandado}, con domicilio procesal en {domicilio_demandado}, a Ud. respetuosamente digo:
            
            Que, vengo a contestar la demanda interpuesta por {demandante}, por los siguientes fundamentos:
            
            I. EXCEPCIONES PRELIMINARES
            {excepciones}
            
            II. DEFENSA DE FONDO
            {defensa}
            
            III. RECONVENCIÓN (si corresponde)
            {reconvencion}
            
            IV. ANEXOS
            1. Copia simple del DNI del demandado.
            2. Poder notarial que acredita la representación.
            3. {anexos}
            
            Por lo expuesto, solicito a Su Señoría:
            1. Tener por contestada la presente demanda.
            2. {petitorio}
            
            Es justicia que solicito.
            
            Lima, {fecha}
            
            _________________________
            {demandado}
            """
        },
        "Admisión": {
            "nombre": "Oficio de Notificación",
            "instrucciones": "Este documento debe ser presentado para notificar al demandado, dentro del plazo de 5 días hábiles.",
            "plantilla": """
            EXPEDIENTE N° {numero_expediente}
            JUZGADO ESPECIALIZADO CIVIL DE LIMA
            {juzgado}
            
            OFICIO DE NOTIFICACIÓN
            
            SEÑOR JUEZ:
            
            Por el presente documento, se notifica a {demandado}, identificado con DNI N° {dni_demandado}, con domicilio en {domicilio_demandado}, sobre la demanda interpuesta por {demandante}.
            
            I. FUNDAMENTOS
            {fundamentos}
            
            II. PLAZO
            Se concede un plazo de 15 días hábiles para contestar la demanda.
            
            III. ANEXOS
            1. Copia de la demanda.
            2. {anexos}
            
            Es justicia que solicito.
            
            Lima, {fecha}
            
            _________________________
            Secretario Judicial
            """
        },
        "Notificación": {
            "nombre": "Contestación de Demanda",
            "instrucciones": "Este documento debe ser presentado en respuesta a la notificación recibida, dentro del plazo de 15 días hábiles.",
            "plantilla": """
            EXPEDIENTE N° {numero_expediente}
            JUZGADO ESPECIALIZADO CIVIL DE LIMA
            {juzgado}
            
            SEÑOR JUEZ:
            
            {demandado}, identificado con DNI N° {dni_demandado}, con domicilio procesal en {domicilio_demandado}, a Ud. respetuosamente digo:
            
            Que, vengo a contestar la demanda interpuesta por {demandante}, por los siguientes fundamentos:
            
            I. EXCEPCIONES PRELIMINARES
            {excepciones}
            
            II. DEFENSA DE FONDO
            {defensa}
            
            III. RECONVENCIÓN (si corresponde)
            {reconvencion}
            
            IV. ANEXOS
            1. Copia simple del DNI del demandado.
            2. Poder notarial que acredita la representación.
            3. {anexos}
            
            Por lo expuesto, solicito a Su Señoría:
            1. Tener por contestada la presente demanda.
            2. {petitorio}
            
            Es justicia que solicito.
            
            Lima, {fecha}
            
            _________________________
            {demandado}
            """
        }
    }
    
    # Obtener el siguiente documento
    siguiente = siguiente_documento.get(etapa_actual)
    if not siguiente:
        return None
    
    # Agregar encabezado con instrucciones
    doc.add_heading(f"PRÓXIMO DOCUMENTO A PRESENTAR: {siguiente['nombre']}", level=1)
    doc.add_paragraph(siguiente['instrucciones'])
    doc.add_paragraph("---" * 20)
    
    # Agregar información del análisis si existe
    if analisis:
        doc.add_heading("INFORMACIÓN RELEVANTE DEL DOCUMENTO ACTUAL", level=2)
        doc.add_paragraph(f"Etapa Detectada: {analisis['etapa']}")
        doc.add_paragraph(f"Nivel de Confianza: {analisis['nivel_confianza'].upper()}")
        
        if analisis['jurisprudencia']:
            doc.add_paragraph("Jurisprudencia Relevante:")
            for jurisprudencia in analisis['jurisprudencia']:
                doc.add_paragraph(f"- {jurisprudencia}")
                doc.add_paragraph(f"  Fundamentos: {obtener_fundamentos_jurisprudencia(jurisprudencia)}")
    
    # Agregar el documento a presentar
    doc.add_heading("DOCUMENTO A PRESENTAR", level=2)
    
    # Preparar datos para la plantilla
    datos_plantilla = {
        "numero_expediente": datos.get("numero_expediente", ""),
        "juzgado": datos.get("juzgado", ""),
        "demandante": datos.get("demandante", ""),
        "dni_demandante": datos.get("dni_demandante", ""),
        "domicilio_demandante": datos.get("domicilio_demandante", ""),
        "demandado": datos.get("demandado", ""),
        "dni_demandado": datos.get("dni_demandado", ""),
        "domicilio_demandado": datos.get("domicilio_demandado", ""),
        "excepciones": "No se formula excepciones preliminares.",
        "defensa": "Se niega la pretensión del demandante por carecer de fundamento legal.",
        "reconvencion": "No se formula reconvención.",
        "anexos": "Documentos de respaldo.",
        "petitorio": "Declarar infundada la demanda.",
        "fecha": datetime.now().strftime("%d de %B del %Y"),
        "fundamentos": "Conforme a la normativa vigente y la jurisprudencia aplicable."
    }
    
    # Formatear la plantilla con los datos
    contenido = siguiente['plantilla'].format(**datos_plantilla)
    doc.add_paragraph(contenido)
    
    # Agregar sección de instrucciones específicas
    doc.add_heading("INSTRUCCIONES ADICIONALES", level=2)
    doc.add_paragraph("1. Presentar el documento en el plazo establecido")
    doc.add_paragraph("2. Adjuntar los documentos de respaldo necesarios")
    doc.add_paragraph("3. Verificar que todos los datos sean correctos")
    doc.add_paragraph("4. Asegurarse de que el documento esté firmado")
    doc.add_paragraph("5. Presentar copias según el número de partes")
    
    # Agregar sección de plazos importantes
    doc.add_heading("PLAZOS IMPORTANTES", level=2)
    plazos = {
        "Contestación de Demanda": "15 días hábiles",
        "Oficio de Notificación": "5 días hábiles",
        "Escrito de Alegatos": "3 días hábiles",
        "Apelación": "5 días hábiles"
    }
    doc.add_paragraph(f"Plazo para presentar: {plazos.get(siguiente['nombre'], 'Verificar en el juzgado')}")
    
    # Agregar sección de requisitos formales
    doc.add_heading("REQUISITOS FORMALES", level=2)
    doc.add_paragraph("1. El documento debe estar en papel A4")
    doc.add_paragraph("2. Debe incluir el número de expediente")
    doc.add_paragraph("3. Debe estar firmado por el abogado y el cliente")
    doc.add_paragraph("4. Debe incluir los anexos correspondientes")
    doc.add_paragraph("5. Debe presentarse en el juzgado correspondiente")
    
    return doc

# --- Integración de imágenes y rediseño visual ---
# Header con logo (balanza de la justicia)
st.markdown("""
    <div style="text-align: center; margin-bottom: 20px;">
        <img src="https://images.unsplash.com/photo-1589829545856-d10d557cf95f?ixlib=rb-1.2.1&auto=format&fit=crop&w=100&q=80" alt="Balanza de la justicia" style="width: 80px; height: 80px; border-radius: 50%;">
        <h1 style="color: #5C4033; font-family: 'Times New Roman', serif; margin-top: 10px;">⚖️ Asistente Legal Peruano ⚖️</h1>
    </div>
""", unsafe_allow_html=True)

# Hero section con imagen de fondo (libros de leyes y mazo de juez)
st.markdown("""
    <div style="background-image: url('https://images.unsplash.com/photo-1589829545856-d10d557cf95f?ixlib=rb-1.2.1&auto=format&fit=crop&w=1200&q=80'); background-size: cover; background-position: center; padding: 50px; border-radius: 10px; margin-bottom: 20px; text-align: center; color: white; text-shadow: 2px 2px 4px rgba(0,0,0,0.5);">
        <h2 style="font-family: 'Times New Roman', serif; font-size: 2.5em;">Tu asistente legal inteligente</h2>
        <p style="font-size: 1.2em;">Analiza documentos, genera escritos y sigue el progreso de tu proceso judicial.</p>
    </div>
""", unsafe_allow_html=True)

# Actualizar las tarjetas de acción con la nueva clase
st.markdown("""
    <div style="display: flex; justify-content: space-around; margin: 20px 0;">
        <div class="action-card">
            <img src="https://images.unsplash.com/photo-1589829545856-d10d557cf95f?ixlib=rb-1.2.1&auto=format&fit=crop&w=50&q=80" alt="Balanza">
            <h3>Analizar Documento</h3>
            <p>Pega tu documento y obtén análisis instantáneo.</p>
        </div>
        <div class="action-card">
            <img src="https://images.unsplash.com/photo-1589829545856-d10d557cf95f?ixlib=rb-1.2.1&auto=format&fit=crop&w=50&q=80" alt="Mazo">
            <h3>Generar Documento</h3>
            <p>Crea escritos legales en segundos.</p>
        </div>
        <div class="action-card">
            <img src="https://images.unsplash.com/photo-1589829545856-d10d557cf95f?ixlib=rb-1.2.1&auto=format&fit=crop&w=50&q=80" alt="Libros">
            <h3>Seguimiento del Proceso</h3>
            <p>Consulta el estado de tu caso en tiempo real.</p>
        </div>
    </div>
""", unsafe_allow_html=True)

# App title with enhanced styling
st.title("⚖️ Asistente Legal Peruano ⚖️")

# Actualizar los títulos de sección
st.markdown('<div class="section-title">Análisis Inicial de Documento</div>', unsafe_allow_html=True)
st.markdown("Pegue su documento para analizar la etapa del proceso y obtener recomendaciones")

# Document input area
documento_inicial = st.text_area(
    "Ingrese el contenido del documento aquí",
    height=200,
    help="Pegue el contenido completo del documento para su análisis"
)

if documento_inicial:
    col1, col2 = st.columns(2)
    
    with col1:
        tipo_doc_inicial = st.selectbox(
            "Tipo de Documento",
            ["Demanda Civil", "Demanda Laboral", "Demanda Penal"],
            key="tipo_doc_inicial"
        )
    
    with col2:
        if st.button("Analizar Documento", key="analizar_inicial"):
            # Mostrar un spinner mientras se realiza el análisis
            with st.spinner("Analizando documento..."):
                # Realizar el análisis con OpenAI
                resultado_analisis = analizar_documento_con_openai(documento_inicial, tipo_doc_inicial)
                
                if resultado_analisis:
                    # Mostrar nivel de confianza
                    confianza_color = {
                        "alto": "#2ecc71",
                        "medio": "#f1c40f",
                        "bajo": "#e74c3c"
                    }
                    
                    st.markdown(f"""
                    <div class="success-box">
                        <h3>Análisis del Documento</h3>
                        <p><strong>Etapa Detectada:</strong> {resultado_analisis["etapa"]}</p>
                        <p><strong>Nivel de Confianza:</strong> <span style="color: {confianza_color[resultado_analisis['nivel_confianza']]}">{resultado_analisis['nivel_confianza'].upper()}</span></p>
                        <p><strong>Artículos Mencionados:</strong> {', '.join(resultado_analisis["articulos"]) if resultado_analisis["articulos"] else "No detectados"}</p>
                        <p><strong>Jurisprudencia Detectada:</strong> {', '.join(resultado_analisis["jurisprudencia"]) if resultado_analisis["jurisprudencia"] else "No detectada"}</p>
                        <p><strong>Campos Faltantes:</strong> {', '.join(resultado_analisis["campos_faltantes"]) if resultado_analisis["campos_faltantes"] else "No se detectaron campos faltantes"}</p>
                        <p><strong>Recomendaciones:</strong></p>
                        <ul>
                            {''.join([f'<li>{rec}</li>' for rec in resultado_analisis["recomendaciones"]])}
                        </ul>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # Extraer datos del documento original
                    datos_extraidos = {
                        "numero_expediente": re.search(r'EXPEDIENTE N°\s*(\d+)', documento_inicial, re.IGNORECASE),
                        "juzgado": re.search(r'JUZGADO[^:]*:\s*([^\n]+)', documento_inicial, re.IGNORECASE),
                        "demandante": re.search(r'([^,]+),\s*identificado con DNI/RUC', documento_inicial, re.IGNORECASE),
                        "dni_demandante": re.search(r'DNI/RUC N°\s*(\d+)', documento_inicial, re.IGNORECASE),
                        "domicilio_demandante": re.search(r'domicilio[^:]*:\s*([^\n]+)', documento_inicial, re.IGNORECASE),
                        "demandado": re.search(r'contra\s*([^,]+),\s*identificado', documento_inicial, re.IGNORECASE),
                        "dni_demandado": re.search(r'DNI/RUC N°\s*(\d+)', documento_inicial, re.IGNORECASE),
                        "domicilio_demandado": re.search(r'domicilio[^:]*:\s*([^\n]+)', documento_inicial, re.IGNORECASE),
                        "hechos": re.search(r'I\.\s*HECHOS\s*(.*?)(?=II\.)', documento_inicial, re.IGNORECASE | re.DOTALL),
                        "fundamentos": re.search(r'II\.\s*FUNDAMENTOS[^:]*:\s*(.*?)(?=III\.)', documento_inicial, re.IGNORECASE | re.DOTALL),
                        "petitorio": re.search(r'III\.\s*PETITORIO[^:]*:\s*(.*?)(?=IV\.|$)', documento_inicial, re.IGNORECASE | re.DOTALL),
                        "fecha": datetime.now().strftime("%d de %B del %Y")
                    }
                    
                    # Limpiar los datos extraídos
                    datos_limpios = {}
                    for key, value in datos_extraidos.items():
                        if key == "fecha":  # La fecha ya es un string formateado
                            datos_limpios[key] = value
                        elif isinstance(value, re.Match):  # Si es un resultado de regex
                            datos_limpios[key] = value.group(1).strip()
                        else:  # Si no se encontró coincidencia
                            datos_limpios[key] = ""
                            
                        # Asegurarse de que no haya None
                        if datos_limpios[key] is None:
                            datos_limpios[key] = ""
                    
                    # Generar el siguiente documento
                    doc = generar_siguiente_documento(
                        resultado_analisis["etapa"],
                        tipo_doc_inicial.split()[1],
                        datos_limpios,
                        documento_original=documento_inicial,
                        analisis=resultado_analisis
                    )
                    
                    if doc:
                        # Guardar el documento en memoria
                        docx = BytesIO()
                        doc.save(docx)
                        docx.seek(0)
                        
                        # Mostrar vista previa del documento
                        st.markdown("### 📝 Próximo Documento a Presentar")
                        st.markdown("""
                        <div class="info-box">
                            <p>Se ha generado el siguiente documento que debe presentar. 
                            Por favor, revise el contenido y descargue el archivo para su uso.</p>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # Botón para descargar el documento
                        st.download_button(
                            label="Descargar Documento para Presentar",
                            data=docx,
                            file_name=f"Proximo_Documento_{datetime.now().strftime('%Y%m%d')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.warning("No se pudo determinar el siguiente documento a presentar.")
                else:
                    st.markdown("""
                    <div class="warning-box">
                        No se pudo realizar el análisis del documento. Por favor, intente nuevamente.
                    </div>
                    """, unsafe_allow_html=True)

# Sidebar for document type selection
st.sidebar.title("Tipo de Documento")
document_type = st.sidebar.selectbox(
    "Seleccione el tipo de documento",
    ["Demanda Civil", "Demanda Laboral", "Demanda Penal", "Escrito de Contestación", "Escrito de Apelación"]
)

# Actualizar la sección de seguimiento en el sidebar
st.sidebar.markdown('<div class="section-title">Seguimiento del Proceso</div>', unsafe_allow_html=True)

# Opción para consultar o ingresar estado
consulta_tipo = st.sidebar.radio(
    "¿Cómo desea actualizar el estado del proceso?",
    ["Consultar por Expediente", "Ingresar Estado Manualmente"]
)

if consulta_tipo == "Consultar por Expediente":
    expediente_consulta = st.sidebar.text_input("Ingrese el número de expediente")
    if expediente_consulta:
        # Aquí se podría implementar la consulta real al sistema judicial
        st.sidebar.info("Funcionalidad de consulta en desarrollo")
else:
    # Selección manual del estado
    tipo_proceso = st.sidebar.selectbox(
        "Tipo de Proceso",
        ["Civil", "Laboral", "Penal"]
    )
    
    etapas = etapas_proceso[tipo_proceso]
    etapa_actual = st.sidebar.selectbox(
        "Etapa Actual",
        etapas,
        index=0
    )
    
    # Mostrar el termómetro de progreso
    current_index = etapas.index(etapa_actual)
    fig = create_process_thermometer(current_index + 1, len(etapas))
    st.sidebar.plotly_chart(fig, use_container_width=True)
    
    # Mostrar próximos pasos
    st.sidebar.markdown("### Próximos Pasos")
    if current_index < len(etapas) - 1:
        st.sidebar.info(f"Siguiente etapa: {etapas[current_index + 1]}")
    else:
        st.sidebar.success("¡Proceso completado!")

# Agregar nueva sección en el sidebar
st.sidebar.markdown("---")
st.sidebar.title("Análisis de Documento")

# Opción para analizar documento existente
analizar_doc = st.sidebar.checkbox("Analizar Documento Existente")

if analizar_doc:
    st.sidebar.markdown("### Ingrese el Documento")
    tipo_doc = st.sidebar.selectbox(
        "Tipo de Documento",
        ["Demanda Civil", "Demanda Laboral", "Demanda Penal"]
    )
    
    documento_texto = st.sidebar.text_area(
        "Pegue el contenido del documento aquí",
        height=200
    )
    
    if documento_texto:
        etapa_detectada = analizar_documento_con_openai(documento_texto, tipo_doc)
        
        if etapa_detectada:
            st.sidebar.success(f"Etapa detectada: {etapa_detectada}")
            
            # Mostrar acciones recomendadas
            st.sidebar.markdown("### Próximas Acciones Recomendadas")
            if tipo_doc in acciones_recomendadas and etapa_detectada in acciones_recomendadas[tipo_doc]:
                for accion in acciones_recomendadas[tipo_doc][etapa_detectada]:
                    st.sidebar.markdown(f"- {accion}")
            
            # Actualizar el termómetro
            etapas = etapas_proceso[tipo_doc.split()[1]]
            try:
                current_index = etapas.index(etapa_detectada)
                fig = create_process_thermometer(current_index + 1, len(etapas))
                st.sidebar.plotly_chart(fig, use_container_width=True)
            except ValueError:
                st.sidebar.warning("No se pudo determinar la posición exacta en el proceso")
        else:
            st.sidebar.warning("No se pudo detectar la etapa del proceso en el documento")

# Main form for case details
with st.form("case_details"):
    st.subheader("Detalles del Caso")
    
    # Common fields
    case_number = st.text_input("Número de Expediente (si existe)")
    court = st.text_input("Juzgado o Sala")
    
    # Plaintiff/Defendant information
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Datos del Demandante")
        plaintiff_name = st.text_input("Nombre del Demandante")
        plaintiff_doc = st.text_input("DNI/RUC del Demandante")
        plaintiff_address = st.text_area("Dirección del Demandante")
    
    with col2:
        st.subheader("Datos del Demandado")
        defendant_name = st.text_input("Nombre del Demandado")
        defendant_doc = st.text_input("DNI/RUC del Demandado")
        defendant_address = st.text_area("Dirección del Demandado")
    
    # Case specific details
    st.subheader("Detalles Específicos del Caso")
    case_description = st.text_area("Descripción del Caso")
    
    # Additional fields based on document type
    if document_type == "Demanda Civil":
        claim_amount = st.number_input("Monto Demandado (S/.)", min_value=0.0)
        legal_basis = st.text_area("Fundamentos de Derecho")
    elif document_type == "Demanda Laboral":
        employment_start = st.date_input("Fecha de Inicio de Relación Laboral")
        employment_end = st.date_input("Fecha de Cese de Relación Laboral")
        last_salary = st.number_input("Último Salario (S/.)", min_value=0.0)
    elif document_type == "Demanda Penal":
        crime_type = st.selectbox("Tipo de Delito", ["Robo", "Hurto", "Estafa", "Otro"])
        crime_date = st.date_input("Fecha del Delito")
        crime_location = st.text_input("Lugar del Delito")
    
    # Submit button
    submitted = st.form_submit_button("Generar Documento")

if submitted:
    st.subheader("Documento Generado")
    
    # Generate document based on type
    if document_type == "Demanda Civil":
        st.markdown(f"""
        ### DEMANDA CIVIL
        
        EXPEDIENTE N° {case_number}
        
        JUZGADO: {court}
        
        SEÑOR JUEZ:
        
        {plaintiff_name}, identificado con DNI/RUC N° {plaintiff_doc}, con domicilio en {plaintiff_address}, a Ud. respetuosamente digo:
        
        Que, vengo a interponer DEMANDA CIVIL contra {defendant_name}, identificado con DNI/RUC N° {defendant_doc}, con domicilio en {defendant_address}, por los siguientes fundamentos:
        
        I. HECHOS
        {case_description}
        
        II. FUNDAMENTOS DE DERECHO
        {legal_basis}
        
        III. PETICIÓN
        Por lo expuesto, solicito a Su Señoría:
        1. Tener por interpuesta la presente demanda.
        2. Ordenar la notificación de la demanda al demandado.
        3. Condenar al demandado al pago de S/. {claim_amount:,.2f}
        
        Es justicia que solicito.
        
        Lima, {datetime.now().strftime('%d de %B del %Y')}
        
        _________________________
        {plaintiff_name}
        """)
    
    elif document_type == "Demanda Laboral":
        st.markdown(f"""
        ### DEMANDA LABORAL
        
        EXPEDIENTE N° {case_number}
        
        JUZGADO: {court}
        
        SEÑOR JUEZ:
        
        {plaintiff_name}, identificado con DNI/RUC N° {plaintiff_doc}, con domicilio en {plaintiff_address}, a Ud. respetuosamente digo:
        
        Que, vengo a interponer DEMANDA LABORAL contra {defendant_name}, identificado con DNI/RUC N° {defendant_doc}, con domicilio en {defendant_address}, por los siguientes fundamentos:
        
        I. HECHOS
        {case_description}
        
        II. RELACIÓN LABORAL
        - Fecha de inicio: {employment_start.strftime('%d/%m/%Y')}
        - Fecha de cese: {employment_end.strftime('%d/%m/%Y')}
        - Último salario: S/. {last_salary:,.2f}
        
        III. PETICIÓN
        Por lo expuesto, solicito a Su Señoría:
        1. Tener por interpuesta la presente demanda laboral.
        2. Ordenar la notificación de la demanda al demandado.
        3. Declarar fundada la demanda y condenar al pago de las prestaciones sociales adeudadas.
        
        Es justicia que solicito.
        
        Lima, {datetime.now().strftime('%d de %B del %Y')}
        
        _________________________
        {plaintiff_name}
        """)
    
    elif document_type == "Demanda Penal":
        st.markdown(f"""
        ### DENUNCIA PENAL
        
        EXPEDIENTE N° {case_number}
        
        FISCALÍA: {court}
        
        SEÑOR FISCAL:
        
        {plaintiff_name}, identificado con DNI/RUC N° {plaintiff_doc}, con domicilio en {plaintiff_address}, a Ud. respetuosamente digo:
        
        Que, vengo a formular DENUNCIA PENAL contra {defendant_name}, identificado con DNI/RUC N° {defendant_doc}, con domicilio en {defendant_address}, por los siguientes fundamentos:
        
        I. HECHOS
        {case_description}
        
        II. DETALLES DEL DELITO
        - Tipo de delito: {crime_type}
        - Fecha del delito: {crime_date.strftime('%d/%m/%Y')}
        - Lugar del delito: {crime_location}
        
        III. PETICIÓN
        Por lo expuesto, solicito a Su Señoría:
        1. Tener por presentada la presente denuncia penal.
        2. Ordenar la investigación correspondiente.
        3. Formular denuncia penal contra el denunciado.
        
        Es justicia que solicito.
        
        Lima, {datetime.now().strftime('%d de %B del %Y')}
        
        _________________________
        {plaintiff_name}
        """)
    
    # Add download button for the generated document
    st.download_button(
        label="Descargar Documento",
        data=st.session_state.get('document_text', ''),
        file_name=f"{document_type.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.txt",
        mime="text/plain"
    ) 